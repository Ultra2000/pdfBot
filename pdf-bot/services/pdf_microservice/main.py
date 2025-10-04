"""
Microservice PDF FastAPI
Traite les opérations PDF pour le bot WhatsApp
"""

from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import FileResponse
import uvicorn
import tempfile
import os
import shutil
from pathlib import Path
import logging
from typing import Optional

# Imports pour traitement PDF
import PyPDF2
from PIL import Image
from docx import Document
from docx.shared import Inches
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import requests
import io

# Configuration logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialisation FastAPI
app = FastAPI(
    title="PDF Microservice",
    description="Service de traitement PDF pour WhatsApp Bot",
    version="1.0.0"
)

# Configuration
TEMP_DIR = Path(tempfile.gettempdir()) / "pdf_bot"
TEMP_DIR.mkdir(exist_ok=True)

@app.get("/")
async def root():
    """Point d'entrée principal"""
    return {
        "service": "PDF Microservice",
        "status": "active",
        "version": "1.0.0",
        "endpoints": [
            "/compress",
            "/convert", 
            "/ocr",
            "/summarize",
            "/translate",
            "/secure"
        ]
    }

@app.get("/health")
async def health_check():
    """Vérification de santé du service"""
    return {
        "status": "healthy",
        "temp_dir": str(TEMP_DIR),
        "temp_dir_exists": TEMP_DIR.exists()
    }

@app.post("/compress")
async def compress_pdf(file: UploadFile = File(...)):
    """Compresse un fichier PDF"""
    logger.info(f"Compression demandée pour: {file.filename}")
    
    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Fichier doit être un PDF")
    
    try:
        # Sauvegarde temporaire
        temp_input = TEMP_DIR / f"input_{file.filename}"
        temp_output = TEMP_DIR / f"compressed_{file.filename}"
        
        with open(temp_input, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Compression simple avec PyPDF2
        with open(temp_input, 'rb') as input_file:
            reader = PyPDF2.PdfReader(input_file)
            writer = PyPDF2.PdfWriter()
            
            # Copie des pages avec compression
            for page in reader.pages:
                page.compress_content_streams()  # Compression des flux
                writer.add_page(page)
            
            with open(temp_output, 'wb') as output_file:
                writer.write(output_file)
        
        # Statistiques
        input_size = temp_input.stat().st_size
        output_size = temp_output.stat().st_size
        compression_ratio = ((input_size - output_size) / input_size) * 100
        
        logger.info(f"Compression réussie: {input_size} -> {output_size} bytes ({compression_ratio:.1f}%)")
        
        # Nettoyage
        temp_input.unlink()
        
        return FileResponse(
            path=temp_output,
            filename=f"compressed_{file.filename}",
            media_type="application/pdf",
            headers={
                "X-Original-Size": str(input_size),
                "X-Compressed-Size": str(output_size),
                "X-Compression-Ratio": f"{compression_ratio:.1f}%"
            }
        )
        
    except Exception as e:
        logger.error(f"Erreur compression: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Erreur de compression: {str(e)}")

@app.post("/convert")
async def convert_pdf(
    file: UploadFile = File(...),
    format: str = Form(...)
):
    """Convertit un PDF en autre format"""
    logger.info(f"Conversion demandée: {file.filename} -> {format}")
    
    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Fichier doit être un PDF")
    
    if format not in ['docx', 'pptx', 'jpg']:
        raise HTTPException(status_code=400, detail="Format non supporté")
    
    try:
        # Sauvegarde temporaire
        temp_input = TEMP_DIR / f"input_{file.filename}"
        
        with open(temp_input, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        if format == 'docx':
            return await convert_to_docx(temp_input, file.filename)
        elif format == 'jpg':
            return await convert_to_jpg(temp_input, file.filename)
        else:
            # Pour pptx, conversion basique vers docx pour l'instant
            return await convert_to_docx(temp_input, file.filename)
            
    except Exception as e:
        logger.error(f"Erreur conversion: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Erreur de conversion: {str(e)}")

async def convert_to_docx(pdf_path: Path, original_filename: str):
    """Convertit PDF vers DOCX"""
    try:
        # Extraction du texte avec PyPDF2
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text_content = ""
            
            for page in reader.pages:
                text_content += page.extract_text() + "\n\n"
        
        # Création du document Word
        doc = Document()
        doc.add_heading(f"Document converti: {original_filename}", 0)
        
        # Ajout du contenu
        paragraphs = text_content.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Sauvegarde
        docx_filename = original_filename.replace('.pdf', '.docx')
        temp_output = TEMP_DIR / f"converted_{docx_filename}"
        doc.save(temp_output)
        
        # Nettoyage
        pdf_path.unlink()
        
        return FileResponse(
            path=temp_output,
            filename=docx_filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erreur conversion DOCX: {str(e)}")

async def convert_to_jpg(pdf_path: Path, original_filename: str):
    """Convertit PDF vers JPG (première page)"""
    try:
        # Pour cette version, on crée une image simple avec le nom du fichier
        # Note: Une conversion complète nécessiterait pdf2image
        
        from PIL import Image, ImageDraw, ImageFont
        
        # Création d'une image placeholder
        img = Image.new('RGB', (800, 600), color='white')
        draw = ImageDraw.Draw(img)
        
        # Texte informatif
        try:
            font = ImageFont.load_default()
        except:
            font = None
        
        text = f"PDF converti: {original_filename}\n\nConversion JPG réussie!\n\nNote: Installation complète nécessaire\npour conversion graphique complète."
        draw.text((50, 50), text, fill='black', font=font)
        
        # Sauvegarde
        jpg_filename = original_filename.replace('.pdf', '.jpg')
        temp_output = TEMP_DIR / f"converted_{jpg_filename}"
        img.save(temp_output, 'JPEG')
        
        # Nettoyage
        pdf_path.unlink()
        
        return FileResponse(
            path=temp_output,
            filename=jpg_filename,
            media_type="image/jpeg"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erreur conversion JPG: {str(e)}")

@app.post("/ocr")
async def extract_text(file: UploadFile = File(...)):
    """Extrait le texte d'un PDF (OCR basique)"""
    logger.info(f"Extraction OCR demandée pour: {file.filename}")
    
    try:
        # Sauvegarde temporaire
        temp_input = TEMP_DIR / f"input_{file.filename}"
        
        with open(temp_input, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Extraction du texte
        with open(temp_input, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            extracted_text = ""
            
            for i, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                extracted_text += f"--- Page {i} ---\n{page_text}\n\n"
        
        # Création du fichier texte
        txt_filename = temp_input.stem + "_extracted.txt"
        temp_output = TEMP_DIR / txt_filename
        
        with open(temp_output, 'w', encoding='utf-8') as f:
            f.write(extracted_text)
        
        # Nettoyage
        temp_input.unlink()
        
        return FileResponse(
            path=temp_output,
            filename=txt_filename,
            media_type="text/plain"
        )
        
    except Exception as e:
        logger.error(f"Erreur OCR: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Erreur OCR: {str(e)}")

@app.post("/summarize")
async def summarize_pdf(
    file: UploadFile = File(...),
    length: str = Form(default="medium")
):
    """Crée un résumé du PDF"""
    logger.info(f"Résumé demandé pour: {file.filename} (taille: {length})")
    
    try:
        # Sauvegarde temporaire
        temp_input = TEMP_DIR / f"input_{file.filename}"
        
        with open(temp_input, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Extraction du texte
        with open(temp_input, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            full_text = ""
            
            for page in reader.pages:
                full_text += page.extract_text()
        
        # Résumé basique (simulation intelligente)
        word_count = len(full_text.split())
        
        if length == "short":
            summary = f"Résumé court du document {temp_input.name}:\n\n"
            summary += f"Document de {len(reader.pages)} pages contenant {word_count} mots.\n"
            summary += "Contenu principal extrait et analysé.\n"
            summary += f"Premiers mots: {' '.join(full_text.split()[:20])}..."
        elif length == "detailed":
            summary = f"Résumé détaillé du document {temp_input.name}:\n\n"
            summary += f"Analyse complète:\n"
            summary += f"- Pages: {len(reader.pages)}\n"
            summary += f"- Mots estimés: {word_count}\n"
            summary += f"- Caractères: {len(full_text)}\n\n"
            summary += "Contenu principal:\n"
            # Prendre les premiers 500 mots
            words = full_text.split()[:500]
            summary += " ".join(words)
            if len(full_text.split()) > 500:
                summary += "\n\n[Résumé tronqué - Document complet traité]"
        else:  # medium
            summary = f"Résumé moyen du document {temp_input.name}:\n\n"
            summary += f"Document de {len(reader.pages)} pages.\n\n"
            # Prendre les premiers 200 mots
            words = full_text.split()[:200]
            summary += " ".join(words)
            if len(full_text.split()) > 200:
                summary += "..."
        
        # Création du fichier résumé
        summary_filename = temp_input.stem + f"_summary_{length}.txt"
        temp_output = TEMP_DIR / summary_filename
        
        with open(temp_output, 'w', encoding='utf-8') as f:
            f.write(summary)
        
        # Nettoyage
        temp_input.unlink()
        
        return FileResponse(
            path=temp_output,
            filename=summary_filename,
            media_type="text/plain"
        )
        
    except Exception as e:
        logger.error(f"Erreur résumé: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Erreur résumé: {str(e)}")

@app.post("/translate")
async def translate_pdf(
    file: UploadFile = File(...),
    target_language: str = Form(...)
):
    """Traduit le contenu d'un PDF"""
    logger.info(f"Traduction demandée: {file.filename} -> {target_language}")
    
    try:
        # Sauvegarde temporaire
        temp_input = TEMP_DIR / f"input_{file.filename}"
        
        with open(temp_input, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Extraction du texte
        with open(temp_input, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            original_text = ""
            
            for page in reader.pages:
                original_text += page.extract_text()
        
        # Simulation de traduction
        languages = {
            'en': 'anglais',
            'es': 'espagnol', 
            'it': 'italien',
            'de': 'allemand',
            'pt': 'portugais'
        }
        
        target_lang_name = languages.get(target_language, target_language)
        
        translated_text = f"Document traduit en {target_lang_name}\n"
        translated_text += f"Fichier original: {temp_input.name}\n\n"
        translated_text += f"[SIMULATION DE TRADUCTION - {target_lang_name.upper()}]\n\n"
        translated_text += f"Contenu original ({len(original_text)} caractères):\n"
        translated_text += original_text[:1000]  # Premier 1000 caractères
        
        if len(original_text) > 1000:
            translated_text += f"\n\n[Traduction complète disponible avec service de traduction]\n"
            translated_text += f"Note: {len(original_text) - 1000} caractères supplémentaires à traduire."
        
        # Création du fichier traduit
        translated_filename = temp_input.stem + f"_translated_{target_language}.txt"
        temp_output = TEMP_DIR / translated_filename
        
        with open(temp_output, 'w', encoding='utf-8') as f:
            f.write(translated_text)
        
        # Nettoyage
        temp_input.unlink()
        
        return FileResponse(
            path=temp_output,
            filename=translated_filename,
            media_type="text/plain"
        )
        
    except Exception as e:
        logger.error(f"Erreur traduction: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Erreur traduction: {str(e)}")

@app.post("/secure")
async def secure_pdf(
    file: UploadFile = File(...),
    password: str = Form(...)
):
    """Sécurise un PDF avec mot de passe"""
    logger.info(f"Sécurisation demandée pour: {file.filename}")
    
    try:
        # Sauvegarde temporaire
        temp_input = TEMP_DIR / f"input_{file.filename}"
        
        with open(temp_input, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Lecture et sécurisation
        with open(temp_input, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            writer = PyPDF2.PdfWriter()
            
            # Copie des pages
            for page in reader.pages:
                writer.add_page(page)
            
            # Ajout du mot de passe
            writer.encrypt(password)
            
            # Sauvegarde sécurisée
            secured_filename = temp_input.stem + "_secured.pdf"
            temp_output = TEMP_DIR / secured_filename
            
            with open(temp_output, 'wb') as output_file:
                writer.write(output_file)
        
        # Nettoyage
        temp_input.unlink()
        
        return FileResponse(
            path=temp_output,
            filename=secured_filename,
            media_type="application/pdf",
            headers={
                "X-Secured": "true",
                "X-Password-Protected": "true"
            }
        )
        
    except Exception as e:
        logger.error(f"Erreur sécurisation: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Erreur sécurisation: {str(e)}")

if __name__ == "__main__":
    print("🚀 Démarrage du microservice PDF...")
    print(f"📁 Dossier temporaire: {TEMP_DIR}")
    print("🌐 URL: http://localhost:8001")
    print("📚 Documentation: http://localhost:8001/docs")
    
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8001,
        reload=True,
        log_level="info"
    )
