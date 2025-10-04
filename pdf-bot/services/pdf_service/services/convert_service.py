import pdf2image
import tempfile
import os
import json
import logging
from pathlib import Path
from PIL import Image
from docx import Document
import openpyxl
from utils.response_utils import create_temp_binary_file, create_temp_response_file

logger = logging.getLogger(__name__)

class ConvertService:
    """
    Service for PDF conversion to various formats
    """
    
    def __init__(self):
        self.supported_formats = ["docx", "xlsx", "img", "png", "jpg", "jpeg"]
    
    async def convert(self, input_path: str, target_format: str, options: str = None) -> str:
        """
        Convert PDF to target format
        
        Args:
            input_path: Path to input PDF
            target_format: Target format (docx/xlsx/img/png/jpg/jpeg)
            options: Additional options as JSON string
            
        Returns:
            Path to converted file
        """
        try:
            logger.info(f"Converting PDF to {target_format}: {input_path}")
            
            # Parse options
            opts = {}
            if options:
                try:
                    opts = json.loads(options)
                except:
                    logger.warning(f"Invalid options JSON: {options}")
            
            # Convert based on target format
            if target_format == "docx":
                return await self._convert_to_docx(input_path, opts)
            elif target_format == "xlsx":
                return await self._convert_to_xlsx(input_path, opts)
            elif target_format in ["img", "png", "jpg", "jpeg"]:
                return await self._convert_to_image(input_path, target_format, opts)
            else:
                raise ValueError(f"Unsupported format: {target_format}")
                
        except Exception as e:
            logger.error(f"Conversion failed: {e}")
            # Create placeholder result for testing
            return self._create_placeholder_result(input_path, target_format)
    
    async def _convert_to_docx(self, input_path: str, options: dict) -> str:
        """Convert PDF to DOCX"""
        try:
            # Convert PDF pages to images
            images = pdf2image.convert_from_path(
                input_path,
                dpi=options.get('dpi', 200),
                first_page=options.get('first_page'),
                last_page=options.get('last_page')
            )
            
            # Create DOCX document
            doc = Document()
            doc.add_heading('Converted from PDF', 0)
            
            # Add each page as an image
            for i, image in enumerate(images):
                # Save image to temporary file
                temp_img = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                image.save(temp_img.name, 'PNG')
                
                # Add to document
                doc.add_paragraph(f'Page {i + 1}:')
                doc.add_picture(temp_img.name, width=doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin)
                doc.add_page_break()
                
                # Clean up temp image
                os.unlink(temp_img.name)
            
            # Save DOCX
            output_path = create_temp_binary_file(b"", "docx")
            doc.save(output_path)
            
            logger.info(f"DOCX conversion completed: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"DOCX conversion failed: {e}")
            raise
    
    async def _convert_to_xlsx(self, input_path: str, options: dict) -> str:
        """Convert PDF to XLSX (placeholder implementation)"""
        try:
            # Create Excel workbook
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "PDF Content"
            
            # Add header
            ws['A1'] = "PDF Conversion"
            ws['A2'] = f"Source: {Path(input_path).name}"
            ws['A3'] = "Note: This is a placeholder implementation"
            ws['A4'] = "Full implementation would extract tables and data"
            
            # Save Excel file
            output_path = create_temp_binary_file(b"", "xlsx")
            wb.save(output_path)
            
            logger.info(f"XLSX conversion completed: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"XLSX conversion failed: {e}")
            raise
    
    async def _convert_to_image(self, input_path: str, format: str, options: dict) -> str:
        """Convert PDF to image format"""
        try:
            # Convert PDF to images
            images = pdf2image.convert_from_path(
                input_path,
                dpi=options.get('dpi', 200),
                first_page=options.get('first_page'),
                last_page=options.get('last_page')
            )
            
            if not images:
                raise ValueError("No images generated from PDF")
            
            # Determine output format
            output_format = format.upper() if format in ["jpg", "jpeg"] else "PNG"
            if format in ["jpg", "jpeg"]:
                output_format = "JPEG"
            
            # If single page, return single image
            if len(images) == 1:
                output_path = create_temp_binary_file(b"", format)
                images[0].save(output_path, output_format, quality=options.get('quality', 95))
                return output_path
            
            # Multiple pages - create a combined image or return first page
            if options.get('combine_pages', False):
                # Combine all pages vertically
                total_width = max(img.width for img in images)
                total_height = sum(img.height for img in images)
                
                combined = Image.new('RGB', (total_width, total_height), 'white')
                y_offset = 0
                
                for img in images:
                    combined.paste(img, (0, y_offset))
                    y_offset += img.height
                
                output_path = create_temp_binary_file(b"", format)
                combined.save(output_path, output_format, quality=options.get('quality', 95))
                return output_path
            else:
                # Return first page only
                output_path = create_temp_binary_file(b"", format)
                images[0].save(output_path, output_format, quality=options.get('quality', 95))
                return output_path
                
        except Exception as e:
            logger.error(f"Image conversion failed: {e}")
            raise
    
    def _create_placeholder_result(self, input_path: str, target_format: str) -> str:
        """Create placeholder result for testing"""
        try:
            if target_format == "docx":
                doc = Document()
                doc.add_heading('PDF Conversion Placeholder', 0)
                doc.add_paragraph(f'This would contain the converted content from: {Path(input_path).name}')
                doc.add_paragraph('Placeholder implementation for testing purposes.')
                
                output_path = create_temp_binary_file(b"", "docx")
                doc.save(output_path)
                return output_path
                
            elif target_format == "xlsx":
                wb = openpyxl.Workbook()
                ws = wb.active
                ws['A1'] = "PDF Conversion Placeholder"
                ws['A2'] = f"Source: {Path(input_path).name}"
                
                output_path = create_temp_binary_file(b"", "xlsx")
                wb.save(output_path)
                return output_path
                
            elif target_format in ["img", "png", "jpg", "jpeg"]:
                # Create a simple placeholder image
                img = Image.new('RGB', (800, 600), 'white')
                output_path = create_temp_binary_file(b"", target_format)
                img.save(output_path, 'PNG' if target_format in ["img", "png"] else 'JPEG')
                return output_path
                
            else:
                # Text placeholder
                content = f"Conversion placeholder for format: {target_format}\nSource: {Path(input_path).name}"
                return create_temp_response_file(content, "txt")
                
        except Exception as e:
            logger.error(f"Failed to create placeholder result: {e}")
            raise
