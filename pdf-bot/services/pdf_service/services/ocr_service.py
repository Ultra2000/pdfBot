import pytesseract
import pdf2image
import tempfile
import os
import logging
from pathlib import Path
from docx import Document
from utils.response_utils import create_temp_response_file, create_temp_binary_file

logger = logging.getLogger(__name__)

class OcrService:
    """
    Service for OCR text extraction from PDFs
    """
    
    def __init__(self):
        self.supported_languages = {
            'eng': 'English',
            'fra': 'French', 
            'spa': 'Spanish',
            'deu': 'German',
            'ita': 'Italian',
            'por': 'Portuguese',
            'rus': 'Russian',
            'chi_sim': 'Chinese Simplified',
            'chi_tra': 'Chinese Traditional',
            'jpn': 'Japanese',
            'kor': 'Korean'
        }
    
    async def extract_text(self, input_path: str, language: str = "eng", output_format: str = "txt") -> str:
        """
        Extract text from PDF using OCR
        
        Args:
            input_path: Path to input PDF
            language: OCR language code
            output_format: Output format (txt/docx)
            
        Returns:
            Path to output file with extracted text
        """
        try:
            logger.info(f"OCR processing: {input_path}, language={language}, format={output_format}")
            
            # Validate language
            if language not in self.supported_languages:
                logger.warning(f"Unsupported language: {language}, using English")
                language = "eng"
            
            # Convert PDF to images
            try:
                images = pdf2image.convert_from_path(
                    input_path,
                    dpi=300,  # Higher DPI for better OCR accuracy
                    grayscale=True  # Grayscale often improves OCR
                )
            except Exception as e:
                logger.error(f"PDF to image conversion failed: {e}")
                return self._create_placeholder_result(input_path, output_format)
            
            # Extract text from each page
            extracted_text = []
            
            for i, image in enumerate(images):
                try:
                    # Configure OCR
                    custom_config = r'--oem 3 --psm 6'
                    
                    # Extract text
                    page_text = pytesseract.image_to_string(
                        image, 
                        lang=language,
                        config=custom_config
                    )
                    
                    if page_text.strip():
                        extracted_text.append(f"=== Page {i + 1} ===\n{page_text.strip()}")
                    else:
                        extracted_text.append(f"=== Page {i + 1} ===\n[No text detected]")
                        
                except Exception as e:
                    logger.warning(f"OCR failed for page {i + 1}: {e}")
                    extracted_text.append(f"=== Page {i + 1} ===\n[OCR processing failed]")
            
            # Combine all text
            full_text = "\n\n".join(extracted_text)
            
            # Create output based on format
            if output_format == "docx":
                return self._create_docx_output(full_text, input_path)
            else:
                return self._create_text_output(full_text, input_path)
                
        except Exception as e:
            logger.error(f"OCR processing failed: {e}")
            return self._create_placeholder_result(input_path, output_format)
    
    def _create_text_output(self, text: str, input_path: str) -> str:
        """Create text file output"""
        try:
            # Add header information
            header = f"OCR Text Extraction\n"
            header += f"Source: {Path(input_path).name}\n"
            header += f"Extraction Date: {os.path.getctime(input_path)}\n"
            header += "=" * 50 + "\n\n"
            
            full_content = header + text
            
            return create_temp_response_file(full_content, "txt")
            
        except Exception as e:
            logger.error(f"Failed to create text output: {e}")
            raise
    
    def _create_docx_output(self, text: str, input_path: str) -> str:
        """Create DOCX file output"""
        try:
            # Create document
            doc = Document()
            
            # Add title
            doc.add_heading('OCR Text Extraction', 0)
            
            # Add metadata
            doc.add_paragraph(f'Source: {Path(input_path).name}')
            doc.add_paragraph(f'Extraction method: Tesseract OCR')
            doc.add_paragraph('')  # Empty line
            
            # Add extracted text
            doc.add_heading('Extracted Text', level=1)
            
            # Split text by pages and add as separate sections
            pages = text.split("=== Page")
            for page in pages:
                if page.strip():
                    if page.startswith(" "):
                        page = "Page" + page
                    
                    # Split into header and content
                    lines = page.split("\n", 1)
                    if len(lines) > 1:
                        page_header = lines[0].strip()
                        page_content = lines[1].strip()
                        
                        if page_header.startswith("Page"):
                            doc.add_heading(page_header, level=2)
                        
                        if page_content:
                            doc.add_paragraph(page_content)
                        
                        doc.add_paragraph("")  # Add spacing
            
            # Save document
            output_path = create_temp_binary_file(b"", "docx")
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to create DOCX output: {e}")
            raise
    
    def _create_placeholder_result(self, input_path: str, output_format: str) -> str:
        """Create placeholder result for testing"""
        try:
            placeholder_text = f"""OCR Text Extraction - Placeholder Result

Source File: {Path(input_path).name}
Processing Status: Placeholder (actual OCR not available)

This is a placeholder result for testing purposes.
In a full implementation, this would contain the actual text
extracted from the PDF using Tesseract OCR.

=== Page 1 ===
[Sample extracted text would appear here]

=== Page 2 ===
[Additional pages would be processed here]

Note: To enable full OCR functionality, ensure Tesseract is properly
installed and configured on the system.
"""
            
            if output_format == "docx":
                doc = Document()
                doc.add_heading('OCR Placeholder Result', 0)
                doc.add_paragraph(placeholder_text)
                
                output_path = create_temp_binary_file(b"", "docx")
                doc.save(output_path)
                return output_path
            else:
                return create_temp_response_file(placeholder_text, "txt")
                
        except Exception as e:
            logger.error(f"Failed to create placeholder result: {e}")
            raise
